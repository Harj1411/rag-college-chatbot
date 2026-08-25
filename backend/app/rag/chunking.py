import os
import re
from typing import List, Dict, Any
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.core.config import settings

def extract_text_from_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF file.
    Returns a list of dicts: [{"page": 1, "text": "..."}]
    """
    pages_data = []
    
    # Try pypdf first
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            text = clean_text(text)
            if text.strip():
                pages_data.append({"page": idx + 1, "text": text})
        if pages_data:
            return pages_data
    except Exception as e:
        pass

    # Fallback to pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for idx, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                text = clean_text(text)
                if text.strip():
                    pages_data.append({"page": idx + 1, "text": text})
    except Exception as e:
        pass

    return pages_data

def extract_text_from_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text from a DOCX/DOC file including paragraphs, tables, and fallback raw text parsing.
    Returns a list of dicts with sections/paragraphs.
    """
    full_text = []

    # 1. Try python-docx parsing (paragraphs + tables)
    try:
        import docx
        doc = docx.Document(file_path)

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                full_text.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if row_cells:
                    full_text.append(" | ".join(row_cells))
    except Exception:
        pass

    # 2. If docx failed or returned empty text (e.g. zipped XML document)
    if not full_text:
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            with zipfile.ZipFile(file_path) as z:
                xml_content = z.read("word/document.xml")
                tree = ET.fromstring(xml_content)
                texts = [node.text for node in tree.iter() if node.text and node.text.strip()]
                if texts:
                    full_text.append(" ".join(texts))
        except Exception:
            pass

    # 3. Final raw binary fallback for old binary .doc files
    if not full_text:
        try:
            with open(file_path, "rb") as f:
                raw_bytes = f.read()
                found_strings = re.findall(rb'[\x20-\x7E]{4,}', raw_bytes)
                valid_lines = []
                for s in found_strings:
                    try:
                        decoded = s.decode('utf-8', errors='ignore').strip()
                        if len(decoded) > 10 and not decoded.startswith(('Root Entry', 'WordDocument', 'Table', 'SummaryInformation')):
                            valid_lines.append(decoded)
                    except Exception:
                        pass
                if valid_lines:
                    full_text.append("\n".join(valid_lines))
        except Exception:
            pass

    joined = "\n\n".join(full_text)
    cleaned = clean_text(joined)
    return [{"page": 1, "text": cleaned}] if cleaned else []

def extract_text_from_txt(file_path: str) -> List[Dict[str, Any]]:
    """Extracts text from plain text/markdown."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    cleaned = clean_text(text)
    return [{"page": 1, "text": cleaned}] if cleaned else []

def clean_text(text: str) -> str:
    """Normalize whitespace and remove non-printable characters."""
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def chunk_document(
    file_path: str,
    doc_id: str,
    filename: str,
    category: str = "general"
) -> List[Dict[str, Any]]:
    """
    Extracts text and breaks it down into chunks (~800 characters, ~100 overlap)
    with attached metadata for vector database indexing.
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        pages = extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        pages = extract_text_from_docx(file_path)
    else:
        pages = extract_text_from_txt(file_path)

    if not pages:
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""]
    )

    chunks = []
    chunk_global_idx = 0

    for page_info in pages:
        page_num = page_info["page"]
        page_text = page_info["text"]
        
        split_texts = splitter.split_text(page_text)
        for piece in split_texts:
            if not piece.strip():
                continue
            chunk_global_idx += 1
            chunks.append({
                "chunk_id": f"{doc_id}_{chunk_global_idx}",
                "doc_id": doc_id,
                "filename": filename,
                "category": category,
                "page": page_num,
                "chunk_index": chunk_global_idx,
                "text": piece.strip()
            })

    return chunks
